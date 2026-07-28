"""Generate the Boss Daddy Pillar & Tier System Plan as an editable .docx on the Desktop.

Reflects the SHIPPED taxonomy (migration 127 + lib/categories.ts + lib/tags.ts).
Authority: docs/pillar-taxonomy.md. Re-run after taxonomy changes to refresh the doc.
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ORANGE = RGBColor(0xE5, 0x5A, 0x1A)
DARK = RGBColor(0x18, 0x18, 0x1B)
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)


def h1(text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = ORANGE
    return p


def h2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = DARK
    return p


def para(text="", italic=False, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    return p


def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.add_run(text)
    return p


def note(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.color.rgb = GREY
    r.font.size = Pt(9.5)
    return p


# ── Title ────────────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
tr = title.add_run("Boss Daddy — Pillar & Tier System")
tr.bold = True
tr.font.size = Pt(22)
tr.font.color.rgb = DARK
sub = doc.add_paragraph()
sr = sub.add_run("Content taxonomy: pillars, subject tiers, and cross-cutting facets")
sr.italic = True
sr.font.color.rgb = GREY
sr.font.size = Pt(11)
note("Reflects the SHIPPED system — migration 127 applied, build green. "
     "Authority: docs/pillar-taxonomy.md. Generated 2026-07-24.")

# ── 1. Tier model ─────────────────────────────────────────────────────────────
h1("1. How the tier model works")
para("Three levels of subject classification, plus a separate set of filter facets — the "
     "standard Wirecutter / REI shape: broad pillars, mid-level subject groups, specific leaf "
     "tags, with facets kept OUT of the subject tree to prevent combinatorial explosion.")
bullet("Tier 1 — PILLAR: the 10 categories in lib/categories.ts. NOT a tag row; a tag points at "
       "its pillar via category_slug.")
bullet("Tier 2 — SUBJECT GROUP: a topic tag with parent_slug = null and category_slug = the pillar.")
bullet("Tier 3 — LEAF: a topic tag with parent_slug = its subject group.")
bullet("CROSS-CUTTING SUBJECT: a topic tag with category_slug = null — belongs to no single pillar.")
bullet("FACETS (separate axis): life-stage, price, use-case, editorial, audience — cut across all pillars.")
para("A piece still has exactly one canonical pillar (its home). The hierarchy organizes the "
     "vocabulary; it does not give a piece two homes.", italic=True)
note("Legend:  [e] = existed before 127   [+] = added in 127   (X) = cross-cutting")

# ── 2. Facets ────────────────────────────────────────────────────────────────
h1("2. Facets (cross-pillar filters — kept separate)")
para("These answer how / who / how much / editorial stance. Never mixed into the subject tree.")
for f, desc in [
    ("life-stage", "pregnancy, newborn, infant, toddler, preschool, school-age, teen"),
    ("price", "under-$25 / $50 / $100 / $250, premium"),
    ("use-case", "travel, daily, occasional, gift-idea, gear-haul (+ functional apparel)"),
    ("editorial", "top-pick, best-value, hidden-gem, boss-approved, mixed-verdict, buyer-beware, better-options, overhyped"),
    ("audience  [+]", "first-time-dad  (room for single-dad, girl-dad, etc. later)"),
]:
    bullet(f"{f} — {desc}")

# ── 3. Pillars ───────────────────────────────────────────────────────────────
h1("3. The Pillars & Subject Tiers")

pillars = [
    ("1. Kids & Family", "kids-family", [
        ("Baby Gear [+]", "strollers [e] · car-seats [e] · baby-carriers [e] · diapering [e] · nursery-gear [e]"),
        ("Baby Sleep [e]", ""),
        ("Feeding [+]", "formula-feeding [e] · breastfeeding [+]"),
        ("Toddler Gear [+]", ""),
        ("Kids Furniture [+]", ""),
        ("Kids Outdoor Gear [+]", ""),
        ("Activities [+]", ""),
        ("Learning [+]", ""),
    ]),
    ("2. Tools & DIY", "tools-diy", [
        ("Power Tools [e]", ""),
        ("Hand Tools [e]", ""),
        ("Home Improvement [e]", ""),
        ("Workshop [e]", "storage-org (X)"),
    ]),
    ("3. Grilling & Cooking", "grilling-cooking", [
        ("Grills & Smokers [+]", ""),
        ("Outdoor Cooking [e]", ""),
        ("Kitchen Tools [e]", "cast-iron [e] · cutlery [+] · utensils [+]"),
        ("Meal Prep [e]", ""),
        ("Recipes [+]", ""),
        ("Coffee [e]", ""),
    ]),
    ("4. Outdoors & Adventure", "outdoors-adventure", [
        ("Camping [e]", ""),
        ("Hiking [e]", ""),
        ("Fishing [e]", ""),
        ("Hunting [e]", ""),
        ("Water Sports [e]", ""),
        ("Outdoor Apparel [+]", ""),
    ]),
    ("5. Tech & Gadgets", "tech-gadgets", [
        ("EDC / Everyday Carry [e]", ""),
        ("Audio Gear [e]", ""),
        ("Wearables [e]", ""),
        ("Charging & Power [+]", ""),
    ]),
    ("6. Vehicles & Garage", "vehicles-garage", [
        ("Automotive [e]", ""),
        ("Truck Gear [e]", ""),
        ("Detailing [e]", ""),
        ("Garage Setup [+]", "storage-org (X)"),
    ]),
    ("7. Health & Wellness", "health-wellness", [
        ("Fitness [e]", "home-gym [e]"),
        ("Mental Health [e]", "mindfulness [e] · self-help [e]"),
        ("Grooming [e]", ""),
        ("Supplements [+]", ""),
        ("Anti-Aging [+]", ""),
    ]),
    ("8. Home & Lifestyle", "home-lifestyle", [
        ("Furniture [+]", ""),
        ("Appliances [+]", ""),
        ("Organization [e]", "storage-org (X)"),
        ("Cleaning [e]", ""),
        ("Yard & Garden [+]", "yard-work [e] · watering [e]"),
        ("Apparel [e]", "general men's / everyday style only (functional gear -> use-case facet)"),
        ("Pet Gear [e]", ""),
    ]),
]

for name, slug, groups in pillars:
    h2(f"{name}   ({slug})")
    for grp, leaves in groups:
        bullet(grp, level=0)
        if leaves:
            bullet(leaves, level=1)

h2("Cross-cutting subjects (category_slug = null)")
bullet("Faith & Doubt (X) — anchors Table Duty AND Health & Wellness")
bullet("storage-org (X) — Workshop, Garage Setup, Organization")
bullet("smart-home (X) — Tech & Gadgets + Home")
bullet("backup-power (X) — camping / vehicles / home")
bullet("safety-first-aid (X) — outdoors / vehicles / home / kids")

# ── 4. Table Duty ────────────────────────────────────────────────────────────
h1("4. Table Duty — 9th Pillar   (table-duty)")
para("A real pillar (not a tenant). Politics, culture, current affairs, and faith — the things a "
     "father refuses to leave unspoken.")
para("Posture: INQUIRY, not proclamation.", bold=True)
bullet("Student-first, not authority-first — 'I'm still learning this too' is the opening posture.")
bullet("Discussion / thought-experiment format, not declarative verdict.")
bullet("Passionate but humble — strong convictions held with the admission he might be wrong.")
bullet("Socratic — often ends on a sharper question rather than a closed answer.")
para("Distinct VOICE REGISTER (shipped):", bold=True)
para("Scoped voice mode in BOSS_DADDY_SYSTEM + brand-guide §1.6 — student-first, no product links. "
     "The default confident reviewer voice is explicitly OFF here. Protector mode still auto-engages "
     "on vulnerability.")

h2("Locked Tier-2 Themes")
for name, desc in [
    ("Meaning & Purpose", "what a life is for — excellence, sacrifice, work, legacy, the standard."),
    ("Faith & Doubt (X)", "belief, uncertainty, how we speak about God with our children. Cross-cutting."),
    ("Culture & Media", "the stories, screens, and ideas we let into the house — and what they form in our kids."),
    ("Citizenship & Duty", "responsibility to family, community, country, next generation; freedom, trade-offs, what we owe."),
    ("Tough Talks", "death, money, sex, failure, power, regret, conflict, screens."),
]:
    p = doc.add_paragraph(style="List Number")
    p.add_run(name + " — ").bold = True
    p.add_run(desc)

# ── 5. Watch Duty ────────────────────────────────────────────────────────────
h1("5. Watch Duty — 10th Pillar   (watch-duty)")
para("The timely counterpart to Table Duty. Table Duty = the ongoing conversation; Watch Duty = "
     "what's happening on our watch right now. Same inquiry voice register.")
para("Operating rules:", bold=True)
bullet("Every piece answers: why does this matter at the family table? (fathers, family, faith, next gen)")
bullet("'Here's what I'm seeing and the questions it raises' — not partisan verdicts.")
bullet("Freshness window: prominent ~2–6 weeks, then archive or grow into a Table Duty essay.")
bullet("High inclusion bar: if it doesn't move fathers, kids, or the future, it doesn't run.")
bullet("May also carry a cross-cutting subject tag (e.g. a car-recall piece surfaces under Vehicles).")

h2("Locked Tier-2 Themes")
for name in ["Family Impact", "Cultural Shifts", "Policy & Power",
             "Technology & Screens", "Safety & Security", "Generational Stakes"]:
    p = doc.add_paragraph(style="List Number")
    p.add_run(name)

# ── 6. Build status ──────────────────────────────────────────────────────────
h1("6. Build status — SHIPPED")
for s in [
    "Migration 127 (parent_slug + category_slug columns, audience group, merch_tags join table, "
    "all tag inserts/re-parenting) — applied clean to prod.",
    "lib/categories.ts — Table Duty + Watch Duty pillars added; 'Tech & EDC' label -> 'Tech & Gadgets'.",
    "lib/tags.ts — buildTagTree() pillar->subject->leaf helper.",
    "Voice register — BOSS_DADDY_SYSTEM + brand-guide §1.6 inquiry-register block.",
    "docs/pillar-taxonomy.md — rewritten to the hierarchical model (the authority).",
    "next build green; db:types regenerated.",
    "TagPicker UI renders the tree — collapsible pillar sections, leaves nested under their "
    "group, Cross-Cutting last, label filter. One component serves reviews/guides/products.",
]:
    bullet(s)

h2("Fast-follows (not built)")
bullet("Watch Duty freshness/archive mechanic (dated prominence -> promote to a Table Duty essay).")

out = r"C:\Users\msb1c\Desktop\BossDaddy-Pillar-Tier-Plan.docx"
doc.save(out)
print("SAVED:", out)
